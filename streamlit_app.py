"""
Streamlit app: "Perfect CV Generator"
- Upload or paste a Job Description (JD)
- Upload an existing CV (DOCX, PDF, TXT)
- Produces a rewritten, ATS-optimised CV tailored to the JD using OpenAI (if API key present)
- Falls back to a simple rule-based generator if no API key provided
- Allows download as DOCX and copyable text
"""

import os
import re
import tempfile
import textwrap
from io import BytesIO

import streamlit as st

# Optional dependencies (make sure they're in requirements)
try:
    import openai
except Exception:
    openai = None

try:
    import docx2txt
except Exception:
    docx2txt = None

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:
    pdf_extract_text = None

try:
    from docx import Document
except Exception:
    Document = None

# ----------------------
# Utilities: File parsing
# ----------------------
def extract_text_from_docx(file_bytes):
    if docx2txt is None:
        raise RuntimeError("docx2txt not installed. See requirements.")
    # docx2txt accepts a path; write temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes.read())
        tmp.flush()
        path = tmp.name
    try:
        text = docx2txt.process(path)
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
    return text

def extract_text_from_pdf(file_bytes):
    if pdf_extract_text is None:
        raise RuntimeError("pdfminer.six not installed. See requirements.")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes.read())
        tmp.flush()
        path = tmp.name
    try:
        text = pdf_extract_text(path)
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
    return text

def extract_text_from_textfile(file_bytes):
    # assume utf-8 text
    raw = file_bytes.read()
    if isinstance(raw, bytes):
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            return raw.decode("latin-1", errors="ignore")
    return raw

def extract_text_from_upload(uploaded_file):
    # uploaded_file is a stream (BytesIO-like) or UploadedFile from Streamlit
    filename = uploaded_file.name.lower()
    uploaded_file.seek(0)
    if filename.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    elif filename.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif filename.endswith(".txt"):
        return extract_text_from_textfile(uploaded_file)
    else:
        # try to read as text anyway
        try:
            return extract_text_from_textfile(uploaded_file)
        except Exception:
            raise RuntimeError("Unsupported file type. Use .docx, .pdf, or .txt")

# ----------------------
# CV parsing (lightweight)
# ----------------------
def simple_cv_parse(cv_text):
    """
    Very light parsing: split into sections using common headers.
    Returns a dict with keys: summary, experience (list), education (list), skills (list), others (text)
    """
    text = cv_text.replace("\r\n", "\n")
    # Normalize repeated spaces and lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Lowercase for section detection only
    lower = text.lower()

    sections = {"summary": "", "experience": [], "education": [], "skills": [], "other": ""}

    # Look for section headers
    # We'll try to split by common headers
    headers = ["experience", "work experience", "employment history", "professional experience",
               "education", "academic qualifications", "qualifications",
               "skills", "technical skills", "key skills",
               "summary", "professional summary", "profile", "objective",
               "projects", "certifications", "achievements"]

    # Build a regex that finds header lines
    header_regex = r"(^|\n)\s*(%s)\s*[:\-]?\s*\n" % "|".join([re.escape(h) for h in headers])
    # find header positions
    splits = [(m.start(), m.group(2).lower()) for m in re.finditer(header_regex, text, re.IGNORECASE | re.MULTILINE)]
    # If no headers, put everything into 'other' and try to detect skills via commas/lines
    if not splits:
        sections["other"] = text.strip()
        # Try to get skills: look for a line with many commas or the word 'skills' not as header
        lines = text.splitlines()
        for line in lines:
            if ',' in line and len(line.split(',')) > 2 and len(sections["skills"]) == 0:
                sections["skills"] = [s.strip() for s in line.split(',') if s.strip()]
                break
        return sections

    # build ranges between headers
    ranges = []
    for i, (pos, hdr) in enumerate(splits):
        start = pos
        end = splits[i+1][0] if i+1 < len(splits) else len(text)
        ranges.append((start, end, hdr))
    # For each range, map to our canonical keys
    for start, end, hdr in ranges:
        chunk = text[start:end].strip()
        # remove the header line
        # first newline after header:
        chunk_body = "\n".join(chunk.splitlines()[1:]).strip()
        if any(h in hdr for h in ["summary", "profile", "objective"]):
            sections["summary"] += "\n" + chunk_body
        elif "experience" in hdr or "employment" in hdr or "professional experience" in hdr:
            # split bullet points or lines into list
            items = [l.strip("•- *") for l in re.split(r"\n(?=\s*[-•*]|\s*[A-Z][a-z]+.*\d{4})", chunk_body) if l.strip()]
            if not items:
                items = [l for l in chunk_body.splitlines() if l.strip()]
            sections["experience"].extend(items)
        elif "education" in hdr or "academic" in hdr or "qualifications" in hdr:
            items = [l for l in chunk_body.splitlines() if l.strip()]
            sections["education"].extend(items)
        elif "skills" in hdr:
            # split by commas or newlines
            raw_skills = re.split(r"[,\n]", chunk_body)
            skills = [s.strip("•- *") for s in raw_skills if s.strip()]
            sections["skills"].extend(skills)
        else:
            sections["other"] += "\n" + chunk_body

    # Clean up: deduplicate and strip
    sections["experience"] = [s.strip() for s in sections["experience"] if s.strip()]
    sections["education"] = [s.strip() for s in sections["education"] if s.strip()]
    sections["skills"] = list(dict.fromkeys([s.strip() for s in sections["skills"] if s.strip()]))
    sections["summary"] = sections["summary"].strip()
    sections["other"] = sections["other"].strip()

    return sections

# ----------------------
# Prompt construction for LLM
# ----------------------
def build_prompt(job_description, cv_text, tone="professional", extra_instructions=""):
    jd = job_description.strip()
    cv = cv_text.strip()
    prompt = f"""
You are an expert CV/resume writer and talent acquisition specialist. Your task is to rewrite and optimise the candidate's CV to be a perfect match for the supplied Job Description. The rewritten CV should be concise, achievements-focused, ATS-friendly (use keywords from the JD), and formatted with clear sections: Name & Contact (use placeholders if missing), Professional Summary (2-4 lines), Key Skills (bullet list), Work Experience (reverse chronological, each role 3-6 bullet achievement lines with metrics where possible), Education, Certifications (if any), and optionally Projects/Awards.

Job Description:
\"\"\"{jd}\"\"\"

Candidate's existing CV content:
\"\"\"{cv}\"\"\"

Requirements:
1. Keep the candidate's factual content — do not invent degrees or employers that didn't exist. You may rephrase, reorganize, and condense information.
2. Use keywords and phrases from the Job Description (exact words where appropriate) so the CV is optimised for ATS.
3. Convert responsibilities into achievement statements (use STAR-ish format), add quantifiable impact when it can be logically inferred from the CV (do NOT make up numbers; approximate only if explicitly allowed - in this template do not invent exact metrics).
4. Keep length to 1-2 pages (concise). Prioritise content relevant to the JD.
5. Output in plain text in the following structured format with section headers (###):
### NAME (use candidate name extracted from CV if present; otherwise write "Candidate Name")
### CONTACT
### PROFESSIONAL SUMMARY
### KEY SKILLS
### WORK EXPERIENCE
- Role | Employer | Dates
  - Achievement 1
  - Achievement 2
### EDUCATION
### CERTIFICATIONS
### PROJECTS / AWARDS (if applicable)

Tone: {tone}
{extra_instructions}

Produce only the final CV text (no commentary). Make it ready-to-paste into Word or ATS systems.
"""
    return textwrap.dedent(prompt)

# ----------------------
# OpenAI interaction
# ----------------------
def call_openai_chat(prompt, model="gpt-4o-mini", max_tokens=1000, temperature=0.2):
    """
    Calls OpenAI ChatCompletion API with the prompt. Requires environment variable OPENAI_API_KEY.
    This function assumes the openai library is installed and configured.
    """
    if openai is None:
        raise RuntimeError("openai package not installed. Install openai in requirements.")
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY not set in environment. Set it or use fallback mode.")
    openai.api_key = api_key

    # Use ChatCompletion if available
    try:
        resp = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert resume writer and recruiter."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=max_tokens,
            temperature=temperature
        )
        # ChatCompletion returns choices with message.content
        content = resp["choices"][0]["message"]["content"]
        return content
    except Exception as e:
        # attempt a fallback to Completion (older clients)
        try:
            resp = openai.Completion.create(
                model="text-davinci-003",
                prompt=prompt,
                max_tokens=max_tokens,
                temperature=temperature,
                n=1
            )
            return resp.choices[0].text
        except Exception as e2:
            raise RuntimeError(f"OpenAI API call failed: {e} / {e2}")

# ----------------------
# Simple rule-based fallback rewrite
# ----------------------
def rule_based_rewrite(job_description, parsed_cv):
    """
    Basic heuristic: pull most relevant skills and experiences, promote them, and produce a formatted CV.
    This is a fallback if no OpenAI key is present.
    """
    jd = job_description.lower()
    # heuristically detect name from top lines
    name = "Candidate Name"
    other_text = parsed_cv.get("other", "")
    first_lines = other_text.splitlines()
    for ln in first_lines[:5]:
        if len(ln.strip()) > 2 and len(ln.split()) <= 4:
            name = ln.strip()
            break

    # pick top skills that appear in JD
    skills = parsed_cv.get("skills", [])
    matched_skills = [s for s in skills if any(word in jd for word in s.lower().split())]
    # if none matched, take first 10
    if not matched_skills:
        matched_skills = skills[:10]

    # build summary: first lines of existing summary or first experience line
    summary = parsed_cv.get("summary") or (parsed_cv.get("experience")[0] if parsed_cv.get("experience") else "")
    summary_short = (summary[:400] + "...") if len(summary) > 400 else summary

    # Build experience: pick top 4 experiences
    exps = parsed_cv.get("experience", [])[:6]
    experience_text = ""
    for i, ex in enumerate(exps):
        # try to extract role and employer if possible
        lines = ex.splitlines()
        header = lines[0]
        bullets = lines[1:] if len(lines) > 1 else []
        experience_text += f"- {header}\n"
        if bullets:
            for b in bullets[:4]:
                experience_text += f"  - {b.strip()}\n"
        else:
            # split by sentences to create bullets
            sents = re.split(r"\. ", ex)
            for s in sents[:3]:
                if s.strip():
                    experience_text += f"  - {s.strip()}\n"
    education = parsed_cv.get("education", [])[:3]

    # Compose final text
    out = []
    out.append(f"### {name}\n")
    out.append("### CONTACT\n- Email: candidate@example.com\n- Phone: +234-000-000-000\n")
    out.append("### PROFESSIONAL SUMMARY\n" + summary_short + "\n")
    out.append("### KEY SKILLS\n" + ("\n".join([f"- {s}" for s in matched_skills]) if matched_skills else "- [Add relevant skills]") + "\n")
    out.append("### WORK EXPERIENCE\n" + experience_text + "\n")
    out.append("### EDUCATION\n" + ("\n".join([f"- {e}" for e in education]) if education else "- [Education details]") + "\n")
    out.append("### CERTIFICATIONS\n- [Add certifications if any]\n")
    out.append("### PROJECTS / AWARDS\n- [Add projects or awards]\n")

    return "\n".join(out)

# ----------------------
# DOCX creation
# ----------------------
def create_docx_from_text(cv_text):
    if Document is None:
        raise RuntimeError("python-docx is not installed. Install via requirements.")
    doc = Document()
    for line in cv_text.splitlines():
        if line.startswith("### "):
            # section header
            doc.add_heading(line.replace("### ", "").strip(), level=2)
        elif line.startswith("- "):
            # bullet
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line.strip())
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ----------------------
# Streamlit UI
# ----------------------
def main():
    st.set_page_config(page_title="Perfect CV Generator", layout="wide")
    st.title("Agentic CV Generator — Turn a CV + Job Description into a tailored, ATS-ready CV")

    st.markdown(
        """
Upload a job description and your current CV. The tool will parse the CV and create a version tailored to the JD.
It uses the OpenAI API if you provide `OPENAI_API_KEY`, otherwise it uses a conservative rule-based fallback.
"""
    )

    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader("1) Job Description")
        jd_text = st.text_area("Paste job description here (or upload a file)", height=220)
        jd_file = st.file_uploader("Or upload a Job Description file (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"])
        if jd_file and not jd_text.strip():
            try:
                jd_text = extract_text_from_upload(jd_file)
                st.success("Job description loaded from file.")
            except Exception as e:
                st.error(f"Could not extract text from JD file: {e}")

    with col2:
        st.subheader("2) Upload your current CV")
        cv_file = st.file_uploader("Upload CV (.docx, .pdf, .txt)", type=["docx", "pdf", "txt"])
        if cv_file is None:
            st.info("If you don't upload a CV, a template skeleton will be generated.")
        st.subheader("Options")
        tone = st.selectbox("Tone for the CV:", options=["professional", "concise", "impactful", "academic"], index=0)
        use_openai = st.checkbox("Use OpenAI API (requires OPENAI_API_KEY set in environment)", value=bool(os.getenv("OPENAI_API_KEY")))
        advanced = st.expander("Advanced options")
        with advanced:
            model = st.text_input("OpenAI model to use", value="gpt-4o-mini")
            max_tokens = st.slider("Max tokens for LLM response", min_value=500, max_value=2500, value=1200, step=100)
            temperature = st.slider("Temperature (creativity)", min_value=0.0, max_value=1.0, value=0.2, step=0.05)

    # parse CV
    raw_cv_text = ""
    if cv_file:
        try:
            raw_cv_text = extract_text_from_upload(cv_file)
        except Exception as e:
            st.error(f"Error extracting CV: {e}")
            raw_cv_text = ""
    else:
        raw_cv_text = ""

    # If user provided no JD, warn
    if not jd_text.strip():
        st.warning("Please paste the job description (or upload a file) to get the best tailored CV.")
    generate_btn = st.button("Generate tailored CV")

    if generate_btn:
        with st.spinner("Parsing CV and generating tailored CV..."):
            if raw_cv_text.strip():
                parsed = simple_cv_parse(raw_cv_text)
            else:
                parsed = {"summary": "", "experience": [], "education": [], "skills": [], "other": ""}

            # build prompt & call OpenAI if requested and available
            final_cv = None
            if use_openai and openai is not None and os.getenv("OPENAI_API_KEY"):
                prompt = build_prompt(jd_text, raw_cv_text or parsed.get("other", ""), tone=tone)
                try:
                    llm_out = call_openai_chat(prompt, model=model, max_tokens=max_tokens, temperature=temperature)
                    final_cv = llm_out.strip()
                except Exception as e:
                    st.error(f"OpenAI call failed: {e}")
                    st.info("Falling back to rule-based rewrite.")
                    final_cv = rule_based_rewrite(jd_text, parsed)
            else:
                # fallback
                final_cv = rule_based_rewrite(jd_text, parsed)

            if final_cv:
                st.success("CV generated.")
                st.subheader("Tailored CV (preview)")
                st.code(final_cv, language="text")
                # Download as DOCX
                if Document is not None:
                    try:
                        docx_buf = create_docx_from_text(final_cv)
                        st.download_button("Download DOCX", data=docx_buf, file_name="tailored_cv.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    except Exception as e:
                        st.warning(f"Could not create DOCX: {e}")
                else:
                    st.info("Install python-docx to enable DOCX download.")
                st.markdown("You can copy the text and paste into your Word/Google Docs or ATS.")
            else:
                st.error("Failed to generate CV.")

    st.markdown("---")
    st.markdown("### Notes and limitations")
    st.markdown(
        """
This tool edits and restructures the input CV using the Job Description. It does not invent factual personal data (dates, employers, degrees). If you use the OpenAI path, do not share extremely sensitive personal data (full DOB, national ID numbers) in the JD or CV text. Always review the generated CV carefully before submitting to employers.
"""
    )

if __name__ == "__main__":
    main()
